
import os
import subprocess
import re

from fastapi import FastAPI, File, UploadFile, Response, status
from pydantic import BaseModel
from typing import List

from PyPDF2 import PdfReader, PdfWriter, PdfFileReader
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import A4
from docx import Document

#from upload_bo import UploadBo

app = FastAPI()

async def makeWatermark():

    # text
    '''pdf = canvas.Canvas("watermark.pdf", pagesize=A4)
    pdf.translate(inch, inch)
    pdf.setFillColor(colors.grey, alpha=0.6)
    pdf.setFont("Helvetica", 50)
    pdf.rotate(45)
    pdf.drawCentredString(400, 100, 'SUPER HIPER Marca')
    pdf.save()'''

    # imagem
    img = ImageReader("./app/logo3.png")
    pdf = canvas.Canvas("watermark.pdf", pagesize=A4)
    pdf.setFillColor(colors.grey, alpha=0.5)
    img_width, img_height = img.getSize()
    aspect = img_height / float(img_width)
    display_width = 380
    display_height = (display_width * aspect)
    pdf.drawImage("./app/logo3.png",
                    (4*cm),
                    ((29.7 - 5) * cm) - display_height,
                    width=display_width,
                    height=display_height,
                    mask="auto")
    pdf.save()


async def addWaterMark(path: str) -> bytes:

    await makeWatermark()

    reader = PdfReader(path)
    page_indices = list(range(0, len(reader.pages)))
    watermark_page = PdfReader("watermark.pdf").pages[0]

    writer = PdfWriter()

    for i in page_indices:
        content_page = reader.pages[i]
        mediabox = content_page.mediabox
        content_page.merge_page(watermark_page)
        content_page.mediabox = mediabox
        writer.add_page(content_page)

    with open(path, "wb") as fp:
        writer.write(fp)

    with open(path, "rb") as arquivo_pdf:
        conteudo_pdf = arquivo_pdf.read()

    return conteudo_pdf


async def convert_to(file: bytes) -> str:
    with open("arquivo.docx", "wb") as arquivo_temporario:
        arquivo_temporario.write(file)

    args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '.', 'arquivo.docx']

    process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=15)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    retorno = await addWaterMark(filename.group(1))

    os.remove(filename.group(1))
    os.remove("arquivo.docx")
    os.remove("watermark.pdf")

    return retorno
       
        
async def replaceString(file: bytes, tags: str, values: str) -> str:
        with open("arquivo.docx", "wb") as arquivo_temporario:
            arquivo_temporario.write(file)

        doc = Document("arquivo.docx")

        tags_split = tags.split(',')
        values_split = values.split(',')

        for section in doc.sections:
            footer = section.footer
            if footer is not None:
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for i in range(len(tags_split)):
                                    if tags_split[i] in paragraph.text:
                                        paragraph.runs[0].text = paragraph.runs[0].text.replace(tags_split[i], values_split[i])
                

        doc.save("arquivo.docx")

        with open("arquivo.docx", "rb") as arquivo_docx:
            conteudo = arquivo_docx.read()

        os.remove("arquivo.docx")

        return conteudo


@app.get('/')
def hello_world():
    return {'message': 'Hello, World!'}

@app.post('/api/docx2pdf', status_code=status.HTTP_200_OK)
async def docx2pdf(file: UploadFile = File(...)):
    arquivo = file.file.read()

    retorno = await convert_to(arquivo)

    return Response(content=retorno,
                    media_type="application/pdf",
                    headers={"Content-Disposition": "attachment; filename=arquivo.pdf"})


@app.post('/api/replace_string')
async def replace_string(tags: str, values: str, file: UploadFile = File(...)):
    arquivo = file.file.read()

    retorno = await replaceString(arquivo, tags, values)
    
    return Response(content=retorno,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": "attachment; filename=arquivo.docx"})
