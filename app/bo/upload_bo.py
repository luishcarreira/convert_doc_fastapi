import os
import subprocess
import re
from typing import List

from PyPDF2 import PdfReader, PdfWriter, PdfFileReader
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import A4

from docx import Document

class UploadBo:

    async def replaceString(file: bytes, tags: str, values: str) -> str:
        with open("arquivo.docx", "wb") as arquivo_temporario:
            arquivo_temporario.write(file)

        doc = Document("arquivo.docx")

        tags_split = tags.split(',')
        values_split = values.split(',')

        print(tags_split)
        print(values_split)

        for section in doc.sections:
            footer = section.footer
            if footer is not None:
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for i in range(len(tags_split)):
                                    if tags[i] in paragraph.text:
                                        paragraph.runs[0].text = paragraph.runs[0].text.replace(tags[i], values[i])
                                # if '#elaborador#' in paragraph.text:
                                #     paragraph.runs[0].text = paragraph.runs[0].text.replace("#elaborador#", "João")
                

        doc.save("arquivo.docx")

        with open("arquivo.docx", "rb") as arquivo_docx:
            conteudo = arquivo_docx.read()

        os.remove("arquivo.docx")

        return conteudo


    async def convertToPdf(file: str):
        #version windows
        '''caminho_docx = "arquivo.docx"

        with open(caminho_docx, "wb") as arquivo_temporario:
            shutil.copyfileobj(file, arquivo_temporario)

        convert(caminho_docx)

        os.remove(caminho_docx)'''


    async def makeWatermark():

        # text
        '''pdf = canvas.Canvas("watermark.pdf", pagesize=A4)
        pdf.translate(inch, inch)
        pdf.setFillColor(colors.grey, alpha=0.6)
        pdf.setFont("Helvetica", 50)
        pdf.rotate(45)
        pdf.drawCentredString(400, 100, 'SUPER HIPER Marca')
        pdf.save()'''

        # imagem
        img = ImageReader('logo.jpg')
        pdf = canvas.Canvas("watermark.pdf", pagesize=A4)
        pdf.setFillColor(colors.grey, alpha=0.6)
        img_width, img_height = img.getSize()
        aspect = img_height / float(img_width)
        display_width = 380
        display_height = (display_width * aspect)
        pdf.drawImage('logo.jpg',
                        (4*cm),
                        ((29.7 - 5) * cm) - display_height,
                        width=display_width,
                        height=display_height,
                        mask="auto")
        pdf.save()


    async def addWaterMark(path: str) -> bytes:

        await UploadBo.makeWatermark()

        reader = PdfReader(path)
        page_indices = list(range(0, len(reader.pages)))

        writer = PdfWriter()

        for i in page_indices:
            content_page = reader.pages[i]

            reader_stamp = PdfReader("watermark.pdf")
            image_page = reader_stamp.pages[0]

            image_page.merge_page(content_page)
            writer.add_page(image_page)

        with open(path, "wb") as fp:
            writer.write(fp)

        with open(path, "rb") as arquivo_pdf:
            conteudo_pdf = arquivo_pdf.read()

        return conteudo_pdf
    

    async def convert_to(file: bytes) -> str:
        with open("arquivo.docx", "wb") as arquivo_temporario:
            arquivo_temporario.write(file)

        args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '.', 'arquivo.docx']
        
        process = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=15)
        filename = re.search('-> (.*?) using filter', process.stdout.decode())
        
        retorno = await UploadBo.addWaterMark(filename.group(1))
 
        os.remove(filename.group(1))
        os.remove("arquivo.docx")
        os.remove("watermark.pdf")

        return retorno